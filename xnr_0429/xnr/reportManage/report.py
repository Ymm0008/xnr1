#!/usr/bin/env python
#encoding: utf-8

from elasticsearch import Elasticsearch
import json
import time
from screenshot import Screen
from docx import Document
from docx.shared import Inches
import xlsxwriter
import sys
sys.path.append('../../')

from parameter import *
reload(sys)
sys.setdefaultencoding('utf-8')

class Report:
	def __init__(self, id_list, username, password,index_name):
		self.ES_HOST = '219.224.134.213'
		self.ES_PORT = 9205
		# self.INDEX_NAME = 'weibo_report_management'
		self.INDEX_NAME = index_name
		self.TYPE = 'report'
		self.es = Elasticsearch([{'host':self.ES_HOST,'port':self.ES_PORT}])
		self.results = []
		self.id_list = id_list
		self.username = username
		self.password = password
		self.currentTime = int(time.time())

	def userList(self):
		self.results = []
		for id in self.id_list:
			result = self.es.get(index=self.INDEX_NAME, doc_type=self.TYPE, id=id)['_source']
			event_name = result['event_name']
			report_time = result['report_time']
			report_type = result['report_type']
			xnr_user_no = result['xnr_user_no']

			weibo_list = json.loads(result['report_content'])['weibo_list']
			for each in weibo_list:
				text = each['text']
				timestamp = each['timestamp']
				try:
					user = each['nick_name']
				except:
					user = each['uid']
				uid = each['uid']
				mid = each['mid']

				dict = {'event_name':event_name, 'report_time':report_time, 'report_type':report_type,\
						'xnr_user_no':xnr_user_no, 'text':text, 'timestamp':timestamp, 'user':user,\
						'uid':uid, 'mid':mid}
				self.results.append(dict)
		return self.results

	def screen_shot(self, results):
		for result in results:
			screen = Screen(self.username, self.password)
			screen.screenShot(result['uid'], result['mid'])

	def save_excel(self):
		results = self.userList()
		filename = 'xnr/static/doc/' + str(self.currentTime)+'.xlsx'
		if results:
			letters = "ABCDEFGHIJKLMN"
			self.screen_shot(results)
			file = xlsxwriter.Workbook(filename)
			table = file.add_worksheet()
			field = [each for each in results[0].keys()]
			field = '^&*'.join(field).replace('event_name',u'上报名称').replace('report_time',u'上报时间').replace('report_type',u'上报类型').replace('xnr_user_no',u'虚拟人').replace('text',u'文本内容').replace('user',u'发博用户').replace('timestamp',u'发博时间')+u"^&*截图"
			field = field.split('^&*')
			for a,b in enumerate(field):
				table.write(letters[a]+str(1),b)
			lists = []
			for result in results:
				list = [each for each in result.values()]
				lists.append(list)
			for i,k in enumerate(lists):
				table.insert_image(letters[len(k)]+str(i+2), results[0]['mid']+'.png', {'x_scale': 0.05, 'y_scale': 0.05})
				for c,d in enumerate(k):
					qq = letters[c]+str(i+2)
					table.write(qq,d)
			file.close()
		return filename

	def save_word(self):
		results = self.userList()
		filename = 'xnr/static/doc/' + str(self.currentTime)+'.docx'
		if results:
			self.screen_shot(results)
			document = Document()
			for result in results:
				result_str = json.dumps(result, ensure_ascii=False).replace('{','').replace('}','').replace('event_name',u'上报名称').replace('report_time',u'上报时间').replace('report_type',u'上报类型').replace('xnr_user_no',u'虚拟人').replace('text',u'文本内容').replace('user',u'发博用户').replace('timestamp',u'发博时间')
				document.add_paragraph(result_str)
				document.add_picture(result['mid']+'.png', width=Inches(1.25))
			document.add_page_break()
			document.save(filename)
		return filename



if __name__ == '__main__':
	report = Report(id_list, SCREEN_WEIBO_USERNAME, SCREEN_WEIBO_PASSWORD,index_name)
	report.save_word()
	report.save_excel()



